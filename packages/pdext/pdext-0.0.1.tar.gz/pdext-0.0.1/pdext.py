import docx
import pandas as pd
import numpy as np


def splitcolumn(df, ColumnName, stripstr, how = 'row'):
    dstripDf = df.drop(ColumnName, 1)
    clist = list(dstripDf.columns)

    if(len(clist) < 2):
        index = df[clist[0]]
    else:
        index = dstripDf

    if (how == 'column'):
        new_df = pd.DataFrame(df[ColumnName].str.split(stripstr).tolist(), index=index)

    elif (how == 'row'):
        new_df = pd.DataFrame(df[ColumnName].str.split(stripstr).tolist(), index=index).stack()

    else:
        return df
    new_df = new_df.reset_index([0, ])
    new_df.rename(columns={new_df.columns[0]: "Oldn_df"}, inplace=True)
    new_df.rename(columns={new_df.columns[1]: ColumnName}, inplace=True)
    if (how == 'row'):
        # If any column value is empty replace it with NaN
        new_df[ColumnName].replace('', np.nan, inplace=True)
        #drop all row with NaN value
        new_df = new_df.dropna(axis=0, how='any', )
    new_df.reset_index(drop=True, inplace=True)

    new_df['Oldn_df'] = new_df['Oldn_df'].map(lambda x: str(x).lstrip('\(').rstrip('\)').replace("\'", ""))

    foo = lambda x: pd.Series([i for i in (x.split(','))])
    InCompleteDf = new_df.Oldn_df.apply(foo)
    InCompleteDf.columns = clist

    #JustValueDF = new_df[ColumnName].to_frame()
    JustValueDF = new_df.drop('Oldn_df', 1)
    CompleteDf = InCompleteDf.join(JustValueDF, how='outer')

    return (CompleteDf)


def addDFtoword(doc, df, addDFHeader=True,  Hbold=True , aftertable='', ifDfEmpty='Empty Table'):
    try:
        doc = doc
    except:
        doc = docx.Document()
    if addDFHeader:
        addDFHeader = 1
    if Hbold:
        Hbold = True
    if df.empty:
        doc.add_paragraph(ifDfEmpty)
    else:
        # df.shape[0] No of rows
        # df.shape[-1] or df.shape[1]  for no of Columns
        table = doc.add_table(df.shape[0] + addDFHeader, df.shape[1], style='Table Grid')  # Extra row for header
        if addDFHeader:
            #add table header as Dataframe header
            for DfHeadercell in range(df.shape[-1]):
                table.cell(0, DfHeadercell).text = df.columns[DfHeadercell]
                table.cell(0, DfHeadercell).paragraphs[0].runs[0].font.bold = Hbold
        # add the rest of the data frame
        for row in range(df.shape[0]):
            for column in range(df.shape[-1]):
                    table.cell(row + addDFHeader, column).text = str(df.values[row, column])

    # Add a row after table with default empty string
    doc.add_paragraph(aftertable)

